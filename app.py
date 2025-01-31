import streamlit as st
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import io

def create_resume():
    st.title("Create Your Resume")
    
    # Personal Information
    st.header("Personal Information")
    name = st.text_input("Full Name", key="name")
    email = st.text_input("Email", key="email")
    phone = st.text_input("Phone", key="phone")
    linkedin = st.text_input("LinkedIn URL", key="linkedin")
    
    # Professional Summary
    st.header("Professional Summary")
    summary = st.text_area("Enter your professional summary", key="summary")
    
    # Projects
    st.header("Projects")
    num_projects = st.number_input("Number of Projects", min_value=0, max_value=10, value=1, key="num_projects")
    projects = []
    
    for i in range(num_projects):
        st.subheader(f"Project {i+1}")
        project_name = st.text_input(f"Project Name {i+1}", key=f"project_name_{i}")
        project_duration = st.text_input(f"Duration {i+1} (e.g., Jun 2021 – July 2021)", key=f"project_duration_{i}")
        project_details = st.text_area(f"Project Details {i+1}", key=f"project_details_{i}")
        projects.append({
            'name': project_name,
            'duration': project_duration,
            'details': project_details.split('\n')
        })
    
    # Experience
    st.header("Work Experience")
    num_experiences = st.number_input("Number of Experiences", min_value=0, max_value=10, value=1, key="num_experiences")
    experiences = []
    
    for i in range(num_experiences):
        st.subheader(f"Experience {i+1}")
        company_name = st.text_input(f"Company Name {i+1}", key=f"company_name_{i}")
        duration = st.text_input(f"Duration {i+1}", key=f"exp_duration_{i}")
        exp_details = st.text_area(f"Experience Details {i+1}", key=f"exp_details_{i}")
        experiences.append({
            'company': company_name,
            'duration': duration,
            'details': exp_details.split('\n')
        })
    
    # Education
    st.header("Education")
    num_education = st.number_input("Number of Education Entries", min_value=0, max_value=5, value=1, key="num_education")
    education = []
    
    for i in range(num_education):
        st.subheader(f"Education {i+1}")
        institution = st.text_input(f"Institution Name {i+1}", key=f"institution_{i}")
        degree = st.text_input(f"Degree/Certificate {i+1}", key=f"degree_{i}")
        edu_duration = st.text_input(f"Duration {i+1}", key=f"edu_duration_{i}")
        grade = st.text_input(f"Grade/Score {i+1}", key=f"grade_{i}")
        education.append({
            'institution': institution,
            'degree': degree,
            'duration': edu_duration,
            'grade': grade
        })
    
    # Skills
    st.header("Skills")
    skills = st.text_area("Enter your skills (one per line)", key="skills")
    
    # Certifications
    st.header("Certifications")
    certifications = st.text_area("Enter your certifications (one per line)", key="certifications")
    
    if st.button("Generate Resume", key="generate_resume"):
        doc = Document()
        
        # Name and Contact Info
        name_paragraph = doc.add_paragraph()
        name_run = name_paragraph.add_run(name)
        name_run.bold = True
        name_run.font.size = Pt(16)
        name_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        contact_paragraph = doc.add_paragraph()
        contact_info = f"{email} | {phone} | {linkedin}"
        contact_run = contact_paragraph.add_run(contact_info)
        contact_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Professional Summary
        doc.add_heading('PROFESSIONAL SUMMARY', level=1)
        doc.add_paragraph(summary)
        
        # Projects
        if projects:
            doc.add_heading('PROJECTS', level=1)
            for project in projects:
                p = doc.add_paragraph()
                p.add_run(f"{project['name']} ({project['duration']})").bold = True
                for detail in project['details']:
                    if detail.strip():
                        doc.add_paragraph(detail.strip(), style='List Bullet')
        
        # Experience
        if experiences:
            doc.add_heading('EXPERIENCE', level=1)
            for exp in experiences:
                p = doc.add_paragraph()
                p.add_run(f"{exp['company']} ({exp['duration']})").bold = True
                for detail in exp['details']:
                    if detail.strip():
                        doc.add_paragraph(detail.strip(), style='List Bullet')
        
        # Education
        if education:
            doc.add_heading('EDUCATION', level=1)
            for edu in education:
                p = doc.add_paragraph()
                p.add_run(f"{edu['institution']} ({edu['duration']})").bold = True
                doc.add_paragraph(f"{edu['degree']} - {edu['grade']}")
        
        # Skills
        if skills:
            doc.add_heading('SKILLS', level=1)
            skills_list = [skill.strip() for skill in skills.split('\n') if skill.strip()]
            doc.add_paragraph(' • '.join(skills_list))
        
        # Certifications
        if certifications:
            doc.add_heading('CERTIFICATION', level=1)
            cert_list = [cert.strip() for cert in certifications.split('\n') if cert.strip()]
            for cert in cert_list:
                doc.add_paragraph(cert, style='List Bullet')
        
        # Save document to bytes buffer
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        
        # Create download button
        st.download_button(
            label="Download Resume",
            data=doc_bytes,
            file_name="resume.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="download_resume"
        )

if __name__ == "__main__":
    create_resume()
